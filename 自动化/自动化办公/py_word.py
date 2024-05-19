
# Python操作Word

# python-docx模块是用于创建和处理Microsoft Word文档的一个Python第三方库，提供了全套的Word操作，是最常用的Word工具
# 官方文档：https://python-docx.readthedocs.io/en/latest/
# 安装：pip install python-docx

# 基本概念：
"""
Document：Word文档对象，多个文档对象互相独立
Paragraph：段落对象，一个Word文档由多个段落组成
Run：节段对象，每个段落由多个节段组成
"""

from docx import Document                  # 用于创建文档
from docx.shared import Inches, Cm, Pt     # 单位
from docx.oxml.ns import qn                # 用于中文字体设置
from docx.shared import RGBColor           # 用于字体颜色设置
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT    # 用于字体背景颜色、段落对齐格式设置
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT               # 用于单元格垂直对齐

# 1、写入Word
# 1）标题：只能设置0-9级标题
# 新建空白文档
doc = Document()
# 设置默认字体、字号和中文字体
style = doc.styles['Normal']
'''
style.font.size = Pt(12)
style.font.name = u'楷体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
'''
# 添加文档标题：add_heading(text, level=1)
title = doc.add_heading('标题', 0)
# 标题居中
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 保存文档
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 2）章节与段落
# 段落
# 创建段落（正文）：add_paragraph(text, style=None)
p1 = doc.add_paragraph("段落1")
# 设置段落两端对齐
p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
# 首行缩进两个字符
p1.paragraph_format.first_line_indent = Cm(0.74)
# 其他缩进：
'''
# 左缩进（英寸）
p1.paragraph_format.left_indent = Inches(0.5)
# 右缩进（磅）
p1.paragraph_format.right_indent = Pt(20)
'''
# 设置行间距：1.5倍行距
p1.paragraph_format.line_spacing = 1.5
# 段前间距（磅）
p1.paragraph_format.space_before = Pt(5)
# 段后间距（磅）
p1.paragraph_format.space_after = Pt(10)
# 创建一级标题
doc.add_heading('一级标题', 1)
# 创建二级标题
doc.add_heading('二级标题', 2)

# 章节（业）
# 创建一个章节
sec = doc.add_section()
'''
# 设置页面高度、宽度
sec.page_height = Inches(15)
sec.page_width = Inches(10)
# 设置页面的边距
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(2)
sec.bottom_margin = Inches(2)
# 设置页眉页脚
head = sec.header
foot = sec.footer
head_par = head.paragraphs[0]
head_par.add_run('页眉')
foot_par = foot.paragraphs[0]
foot_par.add_run('页脚')
'''
# 保存
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 3）字体和引用
# 设置段落文字字体大小（磅）
# 添加文字块：add_run(text, style=None)
run = doc.add_paragraph('段落2：').add_run('设置字号18')
run.font.size = Pt(18)
# 设置英文字体
run = doc.add_paragraph('段落3：').add_run('The Font is Times New Roman')
run.font.name = 'Times New Roman'
# 设置中文字体
run = doc.add_paragraph('段落4：').add_run('设置黑体')
run.font.name = '黑体'
run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
# 设置斜体：run.font.italic = True
doc.add_paragraph('段落5：').add_run('设置斜体').font.italic = True
# 设置粗体：run.font.bold = True
doc.add_paragraph('段落6：').add_run('设置粗体').font.bold = True
# 设置字体带下划线：run.font.underline = True
doc.add_paragraph('段落7：').add_run('设置带下划线').font.underline = True
# 设置字体颜色：run.font.color.rgb
run = doc.add_paragraph('段落8：').add_run('设置字体为红色')
run.font.color.rgb = RGBColor(255, 55, 55)
# 设置字体背景颜色：run.font.highlight_color
run.font.highlight_color = WD_COLOR_INDEX.YELLOW
# 其他设置：
'''
# 设置轮廓线
# run.font.outline = True
# 设置阴影
# run.font.shadow = True
# 设置删除线
# run.font.strike = True
# 设置双删除线
# run.font.double_strike = True
# 设置上标
# run.font.superscript = True
# 设置下标
# run.font.subscript = True
'''
# 添加引用
doc.add_paragraph('添加引用', style='Intense Quote')
# 保存
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 项目列表
# 添加无序列表
doc.add_paragraph('无序列表1', style='List Bullet')
doc.add_paragraph('无序列表2', style='List Bullet')
doc.add_paragraph('无序列表2', style='List Bullet')
# 添加有序列表
doc.add_paragraph('有序列表1', style='List Number')
doc.add_paragraph('有序列表2', style='List Number')
doc.add_paragraph('有序列表3', style='List Number')
# 保存
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 增加分页
doc.add_page_break()
# 设置段落文字在分页时的处理
'''
p1.paragraph_format.keep_together = True        # 段中不分页
p1.paragraph_format.keep_with_next = True       # 与下段同页
p1.paragraph_format.page_break_before = True    # 段前分页
p1.paragraph_format.widow_control = True        # 孤行控制
'''

# 表格：
# 添加表格：add_table(rows, cols, style=None)
# rows:从第几行开始写入  cols:列数  style:样式
table = doc.add_table(rows=1, cols=3, style='Table Grid')
'''
# 指定格式
table.style = 'Table Grid'
# 获取行对象列表
table.rows
# 获取列对象列表
table.columns
# 获取行单元格对象列表
table.row_cells
# 获取列单元格对象列表
table.column_cells
# 增加列：add_column(width)
table.add_column(Inches(3))
# 增加行（只能逐行添加）：add_row()
table.add_row()
'''
# 定义表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '编号'
hdr_cells[1].text = '姓名'
hdr_cells[2].text = '年龄'
# 表格数据
records = [
    (100, 'Tom', 18),
    (101, 'Jerry', 17),
    (102, 'Alice', 28)
]
# 插入数据
for id, name, age in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = name
    row_cells[2].text = str(age)

# 单个值添加数据：table1.cell(0,0).text = '100'
# 补充：合并单元格与对齐
'''
cell_new = table.cell(0, 3).merge(table.cell(1, 3)).merge(table.cell(2, 3))
cell_new.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER       # 垂直居中
cell_new.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER      # 水平居中
'''
# 保存
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 图片：
# 添加图片：add_picture(image_path_or_stream, width=None, height=None)
# 添加空行
doc.add_paragraph('\n')
# 添加图像（Inches：英寸 Cm：厘米）
# 当只设置一个方向的长度（宽或高）时，另一方向会自动缩放
doc.add_picture(r'C:\Users\cc\Desktop\word_test.png', width=Inches(1))
# 保存
doc.save(r'C:\Users\cc\Desktop\test.docx')

# 2、读取Word
# 打开文档
doc = Document(r'C:\Users\cc\Desktop\test.docx')
# 读取标题、段落、列表等内容（除表格）
content = [paragraph.text for paragraph in doc.paragraphs if paragraph != '']
for line in content:
    print(line)

# 读取表格
# 表格取值：cell.text 或 table.cell(row_num, col_num)
tables = [table for table in doc.tables]
tab_list = []
tab = []
for table in tables:
    for row in table.rows:
        row_tup = ()
        cols = len(table.columns)
        for cell in row.cells:
            print(cell.text, end='\t')
            row_tup = row_tup + (cell.text,)
            cols -= 1
            if cols == 0:
                continue
        tab.append(row_tup)
        print()
    tab_list.append(tab)
    print('\n')

print(tab)
print(tab_list)

# 另存为
# doc.save(r'C:\Users\cc\Desktop\test_new.docx')

# 3、将Word表格保存到Excel
from openpyxl import Workbook

workbook = Workbook()
sheet = workbook.active

tab_list = []
tab = []
for table in tables:
    for row in table.rows:
        row_list = []
        cols = len(table.columns)
        for cell in row.cells:
            row_list.append(cell.text)
            cols -= 1
            if cols == 0:
                continue
        tab.append(row_list)
        sheet.append(row_list)
    tab_list.append(tab)

workbook.save(filename=r'C:\Users\cc\Desktop\test.xlsx')
print(tab)
print(tab_list)

# 4、win32com：格式转换（doc转docx）
# 安装：pip install pypiwin32

from win32com import client as wc
from win32com.client import Dispatch, constants

def doc2Docx(doc, docx):
    # 打开Word应用程序
    word = wc.Dispatch('Word.Application')
    # 后台运行，不显示
    word.Visible = 0
    # 不警告
    word.DisplayAlerts = 0
    # 打开doc文件
    doc = word.Documents.Open(doc)
    # 保存为docx文件：txt=4, html=10, docx=12/16， pdf=17
    doc.SaveAs(docx, 12)
    # 关闭Word文档
    doc.Close()
    word.Quit()


# 5、Word转PDF
def word2PDF(doc, pdf):
    # 打开Word应用程序
    word = Dispatch("Word.Application")
    # 后台运行，不显示
    word.Visible = 0
    # 不警告
    word.DisplayAlerts = 0
    # 打开docx文件
    doc = word.Documents.Open(doc)
    # 保存为pdf文件：txt=4, html=10, docx=12/16， pdf=17
    doc.SaveAs(pdf, 17)
    # 关闭Word文档
    doc.Close()
    word.Quit()

# word2PDF(r'C:\Users\cc\Desktop\test.docx', r'C:\Users\cc\Desktop\test.pdf')

